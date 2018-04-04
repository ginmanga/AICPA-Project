""" Modify Wordmaster to parse more general word files
Also needs to read pdfs and determine if they are text or images
"""
import os
import docx
#import string
#import glob


def getText(para, par_top, option = ""):
    """Function gathers text fom docx file"""
    #Will add a variable that takes the list of paragraph numbers within the file
    #And looks for the needed text to gather data
    #doc = docx.Document(filename)
    #Three inputs paras = paragraphs
    #par_top is the list of paragraphs numbers to check
    #options tells how many paragraphs after par_top to gather
    fullText = []
    aicpa_count = 29
    if option is "check":
        aicpa_count = 10
    if option is 'seconline':
        aicpa_count = 40
    for i in par_top:
        newText = []
        for j in range(i, i+aicpa_count):
           (newText.append(para[j].text.strip()) if para[j].text.strip() != '' else None)
        fullText.append(newText)
    return fullText


def parseText(num_docs, text, list_par, doc_type = 'aicpa'):
    """Parse the text gotten from geText"""
    #id_data = ['File Path', 'File Name','Docs in file', 'Company Name', 'SIC', 'DATE', 'TICKER']
    id_data = []
    count = 1
    if doc_type is 'aicpa':
        #names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'start_paragraph', 'Company Name',
                  #'SIC', 'DATE', 'TICKER']]
        months = ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN', 'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC',
                  'MARCH', 'SEPT']
        months.extend([s.strip() + '.' for s in months])
        for i in text:
            name = if_find("COMPANY NAME:", i[1])
            sic = sic_code(i[2:4])
            date = find_strings_lists(i, months)
            ticker = find_strings_lists(i, "Ticker", option=1)
            start_paragraph = list_par[count-1]
            id_data.append([str(count), str(start_paragraph), name, sic, date, ticker])
            count += 1
    if doc_type is 'seconline':
        #names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'start_paragraph',
                  #'Document Type', 'Company Name', 'Filing Date','Document Date',
                  #'TICKER', 'Exchange','Incorporation', 'CUSIP', 'SIC']]
        for i in text:
            doc_type = i[2]
            name = i[4]
            fil_date = i[3].split()[1]
            doc_date = i[3].split()[3]
            ticker = next((if_find('ticker-symbol:', j, 1) for j in i if j.lower().find('ticker-symbol:') >= 0), "").split()[0]
            exchange = next((if_find('ticker-symbol:', j, 1) for j in i if j.lower().find('ticker-symbol:') >= 0), "").split()[2]
            incorp = next((if_find('INCORPORATION:', j, 1) for j in i if j.lower().find('incorporation:') >= 0), "")
            cusip = next((if_find('CUSIP NUMBER:', j, 1) for j in i if j.lower().find('cusip number:') >= 0), "")
            sic = next((if_find('SIC-CODES:', j, 1) for j in i if j.lower().find('sic-codes:') >= 0), "")
            p_sic = next((if_find('PRIMARY SIC:', j, 1) for j in i if j.lower().find('primary sic:') >= 0), "")
            id_data.append([doc_type, name, fil_date, doc_date, ticker, exchange, incorp, cusip, sic, p_sic])
            #print(id_data)
        #print("not done yet")
    return id_data #names

def sic_code(text):
    """Receives raw data and finds the SIC code in AICPA files
    Need to check in which row it is, since some files have an address"""
    a = "SIC CODE:"
    com_sep = str.maketrans("","",";: ") #check how to make this better
    #erases ;: and empty spaces
    try:
        int(if_find(a, text[0]).translate(com_sep))
        code = if_find(a, text[0])
    except:
        code = if_find(a, text[1])
    try:
        int(code.translate(com_sep))
    except:
        code = "NA"
    return code


def if_find(value, text, option = 0):
    """Takes a string and looks for a value
    if found it returns the string without that value (and before)
    else it returns the stripped string"""
    if option == 0:
        if text.find(value) > -1:
            return text[len(value):len(text)].strip()
        else:
            return text.strip()
    if option == 1:
        nt = text.lower().count(value.lower())
        if nt >= 1:
            for i in range(1, nt+1):
                text = text[text.lower().find(value.lower())+len(value):len(text)].strip()
            return text
        else:
            return text.strip()


def find_strings_lists(text, terms, option = 0):
    """Takes two lists of strings. It looks for each element of terms in each element of text
    if it finds it, it gives back the item in text that matches"""
    #option = 0 text and terms are both lists
    if option == 0:
        c1 = next((s for s in text for s1 in terms if s1.lower() in if_find(":",s, option = 1).lower().split()), None)
        try:
            c1 = if_find(":", c1, option = 1)
        except:
            c1 = 'NA'
        return c1
    if option == 1:
        c1 = next((s for s in text if terms.lower() in s.lower()), None)
        try:
            c1 = if_find(terms, c1, option=1)
        except:
            c1 = 'NA'
        return c1


def write_file(path_file, data, options = 0):
    """Writes all data to file"""
    #If no GVKEY or document contains many files, then one CSV file per document
    #For files containing less
    #if we know the gvkey, write a single file with all the files data
    #if no gvkey, write one per file
    path_to_aicpa = os.path.join(path_file, 'summary_aicpa.txt')
    path_to_seconline = os.path.join(path_file, 'summary_seconline.txt')
    path_to_pdf = os.path.join(path_file, 'summary_PDF.txt')
    path_to_fail = os.path.join(path_file, 'summary_FDL.txt')
    #data_ss = open(os.path.join(path_file, 'sum.text'),'w')
    with open(path_to_aicpa,'w') as file:
        file.writelines('\t'.join(i) + '\n' for i in data[0])
    file.close()
    with open(path_to_seconline,'w') as file:
        file.writelines('\t'.join(i) + '\n' for i in data[1])
    with open(path_to_pdf, 'w') as file:
        file.writelines('\t'.join(i) + '\n' for i in data[2])
    with open(path_to_fail, 'w') as file:
        file.writelines('\t'.join(i) + '\n' for i in data[3])
    file.close()


def fnd(paragraphs, terms, terms2, file_name):
    """Given a string of characters find paragraph numbers of each case"""
    #For AICPA files, look for number of number DOCUMENT
    #called by fsttotal
    count_par = 0
    count_doc = 0
    list_paras = []
    aicpa = 0
    for i in paragraphs:
        fc = terms[0] in i.text
        sc = terms[1] in i.text
        dc = any(char.isdigit() for char in i.text)
        sc2 = terms2[0] in i.text
        #dc2 = any(char.isdigit() for char in i.text)
        c_list = [fc, sc, dc]
        c_list2 = [fc, sc2, dc]
        if all(cond is True for cond in c_list) or all(cond is True for cond in c_list2):
            try: #correct for some special cases
                int(i.text.split()[0])
                list_paras.append(count_par)
                count_doc += 1
            except:
                continue
            if count_doc is 1:
                #check if the file is AICPA or SECONLINE
                #text = getText(paragraphs, [count_par], "check")
                doc_type = check_file(paragraphs, [count_par])
                #print(doc_type)
        count_par += 1
    if count_doc is 0:
        print("something wrong with file")
        count_doc, list_paras, doc_type = '0', '', 'FDL'

    return str(count_doc), list_paras, doc_type


def check_file(paragraphs, count_par):
    """Check file type"""
    #call getText and parse first paras after initial document to check type of file
    sec_online = ['copyright', 'sec', 'online']
    text = getText(paragraphs, count_par, "check")
    doc_type = 'aicpa'
    if all(x in text[0][1].lower() for x in sec_online) is True:
        doc_type = "seconline"
    return doc_type


def fsttotal(file_path, file_name): ### NEED TO MODIFY TO INCLUDE AICPA and SEC ONLINE FILES
    """Function to find start and total documents"""
    #Call fnd() function
    los = ['of', 'DOCUMENTS'] #works for most files
    los2 = ['DOCUMENT'] #for special cases
    b = []
    count_doc = ""
    list_paras = ""
    doc_type = ""
    paras = ""
    try:
        file_doc = docx.Document(file_path)
        paras = file_doc.paragraphs
        count_doc, list_paras, doc_type = fnd(paras, los, los2, file_name)
    except docx.opc.exceptions.PackageNotFoundError:
        #print("PDF?")
        doc_type = 'PDF'
        return file_name, count_doc, list_paras, paras, doc_type
    return file_name, count_doc, list_paras, paras, doc_type



def term_gen(type_file): #generate the search terms based on the file
    #First, determine the file type
    if type_file == 'AICPA':
        names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'start_paragraph', 'Company Name',
                  'SIC', 'DATE', 'TICKER']]
        doc_idfer = los = ['of', 'DOCUMENTS']




def file_loop(path, ptofile):
    """This function is called by fipath when the path
    given is a folder and not a file"""
    names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count',
              'start_paragraph', 'Company Name','SIC', 'DATE', 'TICKER']]
    names2 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'start_paragraph',
               'Document Type', 'Company Name', 'Filing Date','Document Date',
               'TICKER', 'Exchange','Incorporation', 'CUSIP', 'SIC', 'Primary SIC']]
    names3 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count']]
    names4 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count']]
    dir_data = []
    #for file in os.listdir(path):
    for file in [f for f in os.listdir(path) if not f.startswith('~$')]:
        #Loops through files and folders in path
        #calls fsttotal function
        file_path_a = os.path.join(path, file)
        if os.path.isdir(file_path_a) is True:
            count = 0
            file_data = [] #resets the data for the new folder
            print([f for f in os.listdir(file_path_a) if not f.startswith('~$')])
            for i in [f for f in os.listdir(file_path_a) if not f.startswith('~$')]:
                file_data = []
                print("HEERE")
                print(i)
                file_path_open = os.path.join(file_path_a, i)
                a = [file_path_open]
                print(a)
                file_name, count_doc, list_paras, paras, doc_type = fsttotal(a[0], os.path.splitext(file)[0])
                a.extend([file_name, count_doc, list_paras])
                if doc_type not in ['PDF', 'FDL']:
                    b = parseText(count_doc, getText(paras, list_paras, doc_type), list_paras,
                                  doc_type)  # collects data from the text in each document
                    for i in b: #append data for files in folder
                        file_data.append(a[0:3]+i)
                else:
                    file_data.append(a)
                count += 1
                for i in file_data:
                    if doc_type == 'aicpa':
                        names.append(i)
                    if doc_type == 'seconline':
                        print("WEEEEE")
                        names2.append(i)
                    if doc_type == 'PDF':
                        names3.append(i)
                    if doc_type == 'FDL':
                        print("INFILEFDL")
                        names4.append(i)
        else:
            file_data = []
            a = [file_path_a]
            file_name, count_doc, list_paras, paras, doc_type = fsttotal(file_path_a, os.path.splitext(file)[0])
            a.extend([file_name, count_doc, list_paras])
            #print(a)
            #print(doc_type)
            if doc_type not in ['PDF', 'FDL']:
                b = parseText(count_doc, getText(paras, list_paras, doc_type), list_paras,
                              doc_type)  # collects data from the text in each document
            #count += 1
                for i in b:  # append data for files in folder
                    file_data.append(a[0:3] + i)
            else:
                file_data.append(a)
            for i in file_data:
                if doc_type == 'aicpa':
                    names.append(i)
                if doc_type == 'seconline':
                    names2.append(i)
                if doc_type == 'PDF':
                    names3.append(i)
                if doc_type == 'FDL':
                    names4.append(i)
    print(names)
    print(names2)
    print(names3)
    print(names4)

    if ptofile == 1:
        write_file(path, [names, names2, names3, names4])
    return names


def fipath(gvkey, path, ptofile = 0):
    """Function delivers path to files to open"""
    # call fsstotal, getText and parseText
    path = os.path.abspath(path)
    names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count',
              'start_paragraph', 'Company Name','SIC', 'DATE', 'TICKER']]
    names2 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'start_paragraph',
               'Document Type', 'Company Name', 'Filing Date','Document Date',
               'TICKER', 'Exchange','Incorporation', 'CUSIP', 'SIC']]
    names3 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count']]
    names4 = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count']]
    if os.path.isdir(path) is False:
        file_name = os.path.splitext(os.path.basename(path))[0]
        # get file name without path or extension
        a = [path]
        file_name, count_doc, list_paras, paras, doc_type = fsttotal(path, file_name)
        a.extend([file_name, count_doc, list_paras])
        file_data = []
        #b = parseText(count_doc, getText(paras, list_paras, doc_type), list_paras, doc_type) #collects data from the text in each document
        #file_data = [a[0:3] + z for z in b]
        print(doc_type)
        if doc_type not in ['PDF', 'FDL']:
            b = parseText(count_doc, getText(paras, list_paras, doc_type), list_paras,
                          doc_type)  # collects data from the text in each document
            # count += 1
            for i in b:  # append data for files in folder
                file_data.append(a[0:3] + i)
        else:
            file_data.append(a)
        for i in file_data:
            if doc_type == 'aicpa':
                names.append(i)
            if doc_type == 'seconline':
                names2.append(i)
            if doc_type == 'PDF':
                names3.append(i)
            if doc_type == 'FDL':
                names4.append(i)
        print(names)
        print(names2)
        print(names3)
        print(names4)
        if ptofile == 1:
            write_file(path, [names, names2, names3, names4])
        return names
    else:
        return file_loop(path, ptofile)
