""" Script to read AICPA Word Files
Not meant to read other types of files
Will have to create other functions for that"""
# First gather identifying data and place it into a spreadsheet
import os
import docx
#import string
import glob


def getText(filename, file_details):
    """Function gathers text fom docx file"""
    #Will add a variable that takes the list of paragraph numbers within the file
    #And looks for the needed text to gather data
    doc = docx.Document(filename)
    fullText = []
    para = doc.paragraphs
    for i in file_details[3]:
        newText = []
        for j in range(i, i+29):
           (newText.append(para[j].text) if para[j].text != '' else None)
        fullText.append(newText)
    return fullText


def parseText(num_docs, text):
    """Parse the text gotten from geText"""
    #id_data = ['File Path', 'File Name','Docs in file', 'Company Name', 'SIC', 'DATE', 'TICKER']
    months = ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN', 'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC',
              'MARCH', 'SEPT']
    months.extend([s.strip()+'.' for s in months])
    id_data = []
    count = 1
    for i in text:
        name = if_find("COMPANY NAME:", i[1])
        sic = sic_code(i[2:4])
        date = find_strings_lists(i, months)
        ticker = find_strings_lists(i, "Ticker", option=1)
        id_data.append([str(count), name, sic, date, ticker])
        count += 1
    return id_data
        #print(ticker)
        #compile data for each file

        #if len(name) >= 64: #special case do not forget
            #print(name)
            #for j in i:
                #print(j)

def sic_code(text):
    """Receives raw data and finds the SIC code in AICPA files
    Need to check in which row it is, since some files have an address"""
    a = "SIC CODE:"
    com_sep = str.maketrans("","",";: ")
    #check how to make this better
    #erases ;: and empty spaces
    try:
        int(if_find(a, text[0]).translate(com_sep))
        code = if_find(a, text[0])
    except:
        code = if_find(a, text[1])
    try:
        int(code.translate(com_sep))
    except:
        code = "NA"
    return code


def if_find(value, text, option = 0):
    """Takes a string and looks for a value
    if found it returns the string without that value (and before)
    else it returns the stripped string"""
    if option == 0:
        if text.find(value) > -1:
            return text[len(value):len(text)].strip()
        else:
            return text.strip()
    if option == 1:
        nt = text.lower().count(value.lower())
        if nt >= 1:
            for i in range(1, nt+1):
                text = text[text.lower().find(value.lower())+len(value):len(text)].strip()
            return text
        else:
            return text.strip()


def find_strings_lists(text, terms, option = 0):
    """Takes two lists of strings. It looks for each element of terms in each element of text
    if it finds it, it gives back the item in text that matches"""
    #option = 0 text and terms are both lists
    if option == 0:
        c1 = next((s for s in text for s1 in terms if s1.lower() in if_find(":",s, option = 1).lower().split()), None)
        try:
            c1 = if_find(":", c1, option = 1)
        except:
            c1 = 'NA'
        return c1
    if option == 1:
        c1 = next((s for s in text if terms.lower() in s.lower()), None)
        try:
            c1 = if_find(terms, c1, option=1)
        except:
            c1 = 'NA'
        return c1


def write_file(path_file, data, options = 0):
    """Writes all data to file"""
    #If no GVKEY or document contains many files, then one CSV file per document
    #For files containing less
    #if we know the gvkey, write a single file with all the files data
    #if no gvkey, write one per file
    print(path_file)
    print(data)
    path_to = os.path.join(path_file, 'summary.txt')
    #data_ss = open(os.path.join(path_file, 'sum.text'),'w')
    #with open(path_to,'w') as file:
        #file.writelines('\t'.join(i) + '\n' for i in data)
    #file.close()

def fnd(paragraphs, terms):
    """Given a string of characters find paragraph numbers of each case"""
    #For AICPA files, look for number of number DOCUMENT
    #called by fsttotal
    count_par = 0
    count_doc = 0
    list_paras = []
    for i in paragraphs:
        fc = terms[0] in i.text
        sc = terms[1] in i.text
        dc = any(char.isdigit() for char in i.text)
        c_list = [fc, sc, dc]
        if all(cond == True for cond in c_list):
            list_paras.append(count_par)
            count_doc += 1
        count_par += 1
    return str(count_doc), list_paras


def fsttotal(file_path, file_name):
    """Function to find start and total documents"""
    #Call fnd() function
    los = ['of', 'DOCUMENTS']
    b = []
    a = [file_name]
    file_doc = docx.Document(file_path)
    paras = file_doc.paragraphs
    a.extend(fnd(paras, los))
    return a


def file_loop(path):
    """This function is called by fipath when the path
    given is a folder and not a file"""
    names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'Company Name', 'SIC', 'DATE', 'TICKER']]
    file_data = []
    for file in os.listdir(path):
        #Loops through files and folders in path
        #calls fsttotal function
        file_path_a = os.path.join(path, file)
        if os.path.isdir(file_path_a) is True:
            count = 0
            for i in os.listdir(file_path_a):
                file_path_open = os.path.join(file_path_a, i)
                a = [file_path_open]
                a.extend(fsttotal(file_path_open, os.path.splitext(i)[0]))
                a.append(getText(file_path_open, a))
                b = parseText(a[2], a[4])  # collects data from the text in each document
                count += 1
                for i in b: #append data when a folder has more than one file
                    file_data.append(a[0:3]+i)
            for i in file_data: #append data of different folders, use it for the ones we know GVKEY
                names.append(i)
            #print(names)
        else:
            a = [file_path_a]
            a.extend(fsttotal(file_path_a, os.path.splitext(file)[0]))
            a.append(getText(file_path_a, a))
            b = parseText(a[2], a[4])
            for i in b:
                file_data.append(a[0:3]+i)
        for i in file_data:
            names.append(i)
        #print(names[1][0])
        #print(names)
    write_file(path, names)
    return names


def fipath(gvkey, path):
    """Function delivers path to files to open"""
    path = os.path.abspath(path)
    if os.path.isdir(path) == False:
        file_name = os.path.splitext(os.path.basename(path))[0]
        # get file name without path or extension
        a = [path]
        a.extend(fsttotal(path, file_name)) #gets starting paragraphs for each document
        a.append(getText(path, a)) #gets initial text of each document
        names = [['File_Path', 'File_Name', 'Doc_num', 'Doc_count', 'Company Name', 'SIC', 'DATE', 'TICKER']]
        b = parseText(a[2], a[4]) #collects data from the text in each document
        file_data = [a[0:3] + z for z in b]
        for i in file_data:
            names.append(i)
        print(names)
        return names
    else:
        return file_loop(path)
